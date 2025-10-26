import os
from langchain_core.documents import Document
import PyPDF2
from docx import Document as DocxDocument  # Alias to avoid conflict

class Reader:
    def __init__(self, filepath):
        self.filepath = filepath

    def read(self):
        """Detects file type and reads content into a LangChain Document."""
        try:
            file_extension = os.path.splitext(self.filepath)[1].lower()
            if file_extension == '.pdf':
                content = self._read_pdf()
            elif file_extension == '.docx':
                content = self._read_word()
            elif file_extension == '.txt':
                content = self._read_text()
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            return Document(page_content=content, metadata={"source": self.filepath})
        except Exception as e:
            raise Exception(f"Error reading file {self.filepath}: {e}")

    def _read_pdf(self):
        """Reads a PDF file and returns its text content."""
        try:
            with open(self.filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {e}")

    def _read_word(self):
        """Reads a Word (docx) file and returns its text content."""
        try:
            doc = DocxDocument(self.filepath)  # Use alias
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX: {e}")

    def _read_text(self):
        """Reads a text file and returns its content."""
        try:
            with open(self.filepath, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error reading TXT: {e}")